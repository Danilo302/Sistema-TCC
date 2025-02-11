from flask import Blueprint, send_file
from app.config import supabase_client
from io import BytesIO
from docx import Document
from docx.shared import Pt
from datetime import date

data_atual = date.today()
data = data_atual.strftime("%d/%m/%Y")


document_bp = Blueprint('preencher-docx', __name__, url_prefix='/preencher-docx')

@document_bp.route('/<uuid:pedido_id>', methods=['POST'])
def preencher_docx(pedido_id):
    
    response = supabase_client.table('pedidos').select("*").eq("id", f"{pedido_id}").execute()
    dados = response.data[0]
    
    id_aluno = dados['id_aluno']
        
        # Consultar informações do aluno
    aluno_response = supabase_client.table('usuarios').select("matricula, nome").eq("id", id_aluno).execute()
    aluno_dados = aluno_response.data[0]

    # Carregar o template DOCX
    doc = Document('app/utils/template.docx')

    # Função para substituir placeholders
    def substituir_placeholder(doc, placeholder, valor):
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, valor)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)  

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, valor)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = "Arial"
                                run.font.size = Pt(11)

    # Substituir placeholders pelos valores recebidos
    substituir_placeholder(doc, '{{semestre}}', dados.get('semestre', ''))
    substituir_placeholder(doc, '{{data_defesa}}', dados.get('data_defesa', ''))
    substituir_placeholder(doc, '{{horario}}', dados.get('horario', ''))
    substituir_placeholder(doc, '{{aluno}}', aluno_dados.get('nome', ''))
    substituir_placeholder(doc, '{{curso}}', dados.get('curso', ''))
    substituir_placeholder(doc, '{{matricula}}',str( aluno_dados.get('matricula', '')))
    substituir_placeholder(doc, '{{turno}}', dados.get('turno', ''))
    substituir_placeholder(doc, '{{orientador}}', dados.get('orientador', ''))
    substituir_placeholder(doc, '{{tema_pesquisa}}', dados.get('tema_pesquisa', ''))
    substituir_placeholder(doc, '{{titulo_provisorio}}', dados.get('titulo', ''))
    substituir_placeholder(doc, '{{atividades_desenvolvidas}}', dados.get('atividades_desenvolvidas', ''))
    substituir_placeholder(doc, '{{contribuicoes_pesquisa}}', dados.get('contribuicoes_pesquisa', ''))
    substituir_placeholder(doc, '{{membro1}}', dados.get('membro1', ''))
    substituir_placeholder(doc, '{{membro2}}', dados.get('membro2', ''))
    substituir_placeholder(doc, '{{data}}', data)

    output = BytesIO()
    doc.save(output)
    output.seek(0)  # Garantir que o ponteiro esteja no início do arquivo

    # Retornar o arquivo gerado para o usuário para download
    return send_file(output, as_attachment=True, download_name="documento_preenchido.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")